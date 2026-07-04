"""
Zynex Export Engine
Generates PDF reports, HTML slide decks, and DOCX documents from ResearchReport + chart images.
"""

from __future__ import annotations

import logging
from io import BytesIO

from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.config import settings
from backend.models.schemas import ResearchReport

logger = logging.getLogger("zynex.agents.export_engine")

# ── Jinja environment ─────────────────────────────────────────────────────────
_jinja_env = Environment(
    loader=FileSystemLoader(str(settings.TEMPLATES_DIR)),
    autoescape=True,
)


def _render_template(
    template_name: str,
    report: ResearchReport,
    charts: dict[str, str],
) -> str:
    """Render a Jinja2 template with the report data."""
    template = _jinja_env.get_template(template_name)
    return template.render(
        report=report,
        charts=charts,
        brand={
            "name": "Zynex",
            "primary_dark": "#0A0A0F",
            "primary_blue": "#3B82F6",
            "primary_cyan": "#06B6D4",
            "accent_violet": "#8B5CF6",
        },
    )


async def generate_pdf(
    report: ResearchReport,
    charts: dict[str, str],
) -> bytes:
    """
    Generate a professional PDF from the report.

    Uses Jinja2 template + WeasyPrint for HTML→PDF conversion.
    Falls back to a simple HTML-based PDF if WeasyPrint is unavailable.
    """
    logger.info("Generating PDF for '%s'", report.topic)

    html_content = _render_template("report.html", report, charts)

    try:
        from weasyprint import HTML
        pdf_bytes = HTML(string=html_content).write_pdf()
        logger.info("PDF generated successfully (%d bytes)", len(pdf_bytes))
        return pdf_bytes
    except ImportError:
        logger.warning(
            "WeasyPrint not installed – returning HTML as PDF fallback. "
            "Install with: pip install weasyprint"
        )
        # Return the styled HTML as bytes (browsers can still render it)
        return html_content.encode("utf-8")
    except Exception as exc:
        logger.error("PDF generation failed: %s", exc)
        # Return the HTML as fallback
        return html_content.encode("utf-8")


async def generate_slides(
    report: ResearchReport,
    charts: dict[str, str],
) -> str:
    """
    Generate an HTML slide deck from the report.

    Returns complete HTML string with embedded CSS/JS for keyboard navigation.
    """
    logger.info("Generating slides for '%s'", report.topic)

    html_content = _render_template("slides.html", report, charts)
    logger.info("Slides generated successfully")
    return html_content


async def generate_docx(
    report: ResearchReport,
    charts: dict[str, str],
) -> bytes:
    """
    Generate a DOCX document from the report.

    Returns DOCX file as bytes.
    """
    logger.info("Generating DOCX for '%s'", report.topic)

    doc = Document()
    
    # Title
    title = doc.add_heading(report.topic, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {report.generated_at}")
    doc.add_paragraph(f"Word Count: {report.word_count}")
    doc.add_paragraph()
    
    # Summary
    if report.summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.summary)
        doc.add_paragraph()
    
    # Sections
    for idx, section in enumerate(report.sections, 1):
        doc.add_heading(section.title or f"Section {idx}", level=1)
        doc.add_paragraph(section.content or "")
        doc.add_paragraph()
    
    # Citations
    if report.citations:
        doc.add_heading("Sources", level=1)
        for idx, citation in enumerate(report.citations, 1):
            p = doc.add_paragraph(f"[{idx}] ", style="List Number")
            p.add_run(citation.title or "Untitled").bold = True
            p.add_run(f"\n{citation.url or ''}")
            doc.add_paragraph()
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("DOCX generated successfully")
    return buffer.getvalue()
